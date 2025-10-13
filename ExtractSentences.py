from docx import Document
import re
import redis
import json

class ExtractSentences:
    def __init__(self, filename):
        self.word_index = {}
        self.doc = Document(filename)
        self.new_doc = Document()
        self.redis_client = redis.Redis(host='localhost', port=6379, decode_responses=True)
        # self.redis_client.flushdb()

    def create_reverse_index(self):
        for i in range(0, len(self.doc.paragraphs)):
            line = self.doc.paragraphs[i].text

            line = re.sub(r'[^\w\s]', '', line)
            line_list = [word.lower() for word in line.split()]

            for word in line_list:
                # if self.word_index.get(word):
                #     self.word_index[word].append(i)
                # else:
                #     self.word_index[word] = [i]

                existing_value = self.redis_client.get(word)
                if existing_value:
                    try:
                        existing_list = json.loads(existing_value)
                        if not isinstance(existing_list, list):
                            existing_list = [existing_list]
                    except (json.JSONDecodeError, TypeError):
                        try:
                            existing_list = [int(existing_value)]
                        except ValueError:
                            existing_list = []

                    existing_list.append(i)
                    self.redis_client.set(word, json.dumps(existing_list))
                else:
                    self.redis_client.set(word, json.dumps([i]))

    def create_new_doc(self, target_word):
        self.new_doc.add_heading(f"{target_word}")

        target_word = target_word.lower()
        # line_numbers = self.word_index.get(target_word)
        line_numbers_json = self.redis_client.get(target_word)

        if line_numbers_json:
            line_numbers = json.loads(line_numbers_json)
            for line_number in line_numbers:
                self.new_doc.add_paragraph(self.doc.paragraphs[line_number].text)

    def save_new_doc(self, target_word):
        self.new_doc.save(f"{target_word}.docx")
